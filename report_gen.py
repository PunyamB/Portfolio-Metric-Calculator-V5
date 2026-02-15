from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GREEN = RGBColor(0x00, 0x7A, 0x33)
RED = RGBColor(0xCC, 0x00, 0x00)
GRAY = RGBColor(0x88, 0x88, 0x88)
LIGHT_GRAY = RGBColor(0xAA, 0xAA, 0xAA)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_cell_text(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    """Add formatted text to a table cell."""
    cell.paragraphs[0].alignment = align
    run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(text)
    if not cell.paragraphs[0].runs or run.text != text:
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # cell padding
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = parse_xml(f'<w:tcMar {nsdecls("w")}>'
                    f'<w:top w:w="40" w:type="dxa"/>'
                    f'<w:bottom w:w="40" w:type="dxa"/>'
                    f'<w:left w:w="80" w:type="dxa"/>'
                    f'<w:right w:w="80" w:type="dxa"/>'
                    f'</w:tcMar>')
    tc_pr.append(mar)


def _add_header_row(table, headers):
    """Format the first row as a header row."""
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        _set_cell_shading(cell, "1B2A4A")
        _add_cell_text(cell, h, bold=True, size=9, color=WHITE)


def _fmt_dollar(v):
    if v is None or v == "N/A":
        return "N/A"
    n = float(v)
    prefix = "-$" if n < 0 else "$"
    return prefix + f"{abs(n):,.2f}"


def _fmt_num(v, decimals=4):
    if v is None or v == "N/A":
        return "N/A"
    return f"{float(v):.{decimals}f}"


def _section_title(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(16)
    p.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = NAVY


def _body_text(doc, text):
    p = doc.add_paragraph()
    p.space_after = Pt(6)
    run = p.add_run(text or "[No narrative generated]")
    run.font.name = "Arial"
    run.font.size = Pt(10)


def _centered_text(doc, text, size=10, color=None, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _make_mover_table(doc, movers, color):
    if not movers:
        p = doc.add_paragraph()
        run = p.add_run("None")
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        return

    headers = ["Ticker", "Company", "Return %", "$ Impact", "Shares", "Status"]
    table = doc.add_table(rows=1 + len(movers), cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_header_row(table, headers)

    for i, m in enumerate(movers):
        row = table.rows[i + 1]
        vals = [
            m.get("ticker", ""),
            m.get("company", m.get("ticker", "")),
            f"{m.get('return_pct', 0):+.2f}%",
            _fmt_dollar(m.get("pnl", 0)),
            str(round(m.get("shares", 0))),
            m.get("note", ""),
        ]
        aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
        colors = [None, None, color, color, None, None]
        bolds = [True, False, False, False, False, False]

        for j, val in enumerate(vals):
            _add_cell_text(row.cells[j], val, bold=bolds[j], color=colors[j], align=aligns[j])


def generate_report(report_data, narrative, output_path="portfolio_report.docx"):
    """Generate a DOCX report from data and narrative."""
    d = report_data
    n = narrative

    doc = Document()

    # default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)

    # margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    _centered_text(doc, f"{d.get('report_type', 'Portfolio')} Portfolio Report", size=20, bold=True, color=NAVY)
    _centered_text(doc, f"{d['period_start']} \u2014 {d['period_end']}", size=11, color=GRAY)
    _centered_text(doc, "Bulls Before Bros | UConn MSFERM FNCE 5322", size=10, color=GRAY)
    doc.add_paragraph()

    # 1. Executive Summary
    _section_title(doc, "1. Executive Summary")
    _body_text(doc, n.get("executive_summary", ""))

    # 2. Performance Snapshot
    _section_title(doc, "2. Performance Snapshot")

    snap_data = [
        ("Portfolio Value", _fmt_dollar(d.get("start_value")), _fmt_dollar(d.get("end_value")),
         _fmt_dollar(d.get("value_change")), d.get("value_change", 0) >= 0),
        ("Positions", str(d.get("start_holdings_count", 0)), str(d.get("end_holdings_count", 0)),
         str(d.get("end_holdings_count", 0) - d.get("start_holdings_count", 0)), True),
        ("Trades in Period", "-", str(d.get("mid_period_trades", 0)), "-", True),
    ]

    metric_keys = ["Sharpe Ratio", "Sortino Ratio", "Beta", "Alpha (Annual)",
                   "VaR (95%)", "Max Drawdown", "Std Dev (Annual)"]
    for key in metric_keys:
        sv = d.get("start_metrics", {}).get(key)
        ev = d.get("end_metrics", {}).get(key)
        delta = d.get("metric_deltas", {}).get(key)
        snap_data.append((
            key,
            _fmt_num(sv) if sv is not None else "N/A",
            _fmt_num(ev) if ev is not None else "N/A",
            (f"{delta:+.4f}" if delta is not None else "N/A"),
            (delta or 0) >= 0,
        ))

    headers = ["Metric", "Start", "End", "Change", "Signal"]
    table = doc.add_table(rows=1 + len(snap_data), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_header_row(table, headers)

    for i, (label, start, end, change, up) in enumerate(snap_data):
        row = table.rows[i + 1]
        _add_cell_text(row.cells[0], label, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        _add_cell_text(row.cells[1], start)
        _add_cell_text(row.cells[2], end)
        ch_color = GREEN if up else RED
        _add_cell_text(row.cells[3], change, color=ch_color)
        _add_cell_text(row.cells[4], "\u2191" if up else "\u2193", color=ch_color)

    # benchmark note
    rf_pct = (d.get("risk_free_rate", 0) * 100)
    p = doc.add_paragraph()
    run = p.add_run(f"Benchmark: S&P 500 return: {d.get('benchmark_return', 0)}% | Risk-free rate: {rf_pct:.2f}%")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    run.font.italic = True

    # 3. Performance Attribution
    _section_title(doc, "3. Performance Attribution")
    _body_text(doc, n.get("performance_attribution", ""))

    p = doc.add_paragraph()
    run = p.add_run("Top Gainers")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = GREEN
    _make_mover_table(doc, d.get("top_gainers", []), GREEN)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Top Losers")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RED
    _make_mover_table(doc, d.get("top_losers", []), RED)

    # 4. Risk Profile
    _section_title(doc, "4. Risk Profile Analysis")
    _body_text(doc, n.get("risk_analysis", ""))

    # 5. Factor Exposures
    fe = d.get("factor_exposures", {})
    factor_names = ["Mkt-RF", "SMB", "HML", "RMW", "CMA"]
    valid_factors = [f for f in factor_names if f in fe]

    if valid_factors:
        _section_title(doc, "5. Factor Exposures")
        fheaders = ["Factor", "Exposure", "t-stat", "Significant", "Interpretation"]
        ftable = doc.add_table(rows=1 + len(valid_factors), cols=5)
        ftable.alignment = WD_TABLE_ALIGNMENT.CENTER
        _add_header_row(ftable, fheaders)

        for i, fname in enumerate(valid_factors):
            fdata = fe[fname]
            row = ftable.rows[i + 1]
            _add_cell_text(row.cells[0], fname, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
            _add_cell_text(row.cells[1], _fmt_num(fdata.get("exposure"), 4))
            _add_cell_text(row.cells[2], _fmt_num(fdata.get("t_stat"), 2))
            _add_cell_text(row.cells[3], "Yes" if fdata.get("significant") else "No")
            interp = str(fdata.get("interpretation", ""))[:45]
            _add_cell_text(row.cells[4], interp, align=WD_ALIGN_PARAGRAPH.LEFT)

    # 6. Outlook
    _section_title(doc, "6. Outlook & Recommendations")
    _body_text(doc, n.get("outlook", ""))

    # Footer
    doc.add_paragraph()
    _centered_text(doc, "\u2014 End of Report \u2014", size=9, color=LIGHT_GRAY)
    _centered_text(doc, "Generated by Paper Trading Portfolio Dashboard | Data: Yahoo Finance, Kenneth French Data Library",
                   size=8, color=LIGHT_GRAY)

    doc.save(output_path)
    return output_path
